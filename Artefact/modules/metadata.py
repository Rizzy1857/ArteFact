"""
Metadata Extraction Module
==========================

Extracts metadata and timestamps from various file types including
images, documents, and media files.
"""

import logging
import subprocess
import json
import struct
from pathlib import Path
from datetime import datetime, timezone
from typing import Dict, List, Optional, Any, Union, Tuple
from rich.console import Console
from rich.table import Table
from rich.progress import Progress, BarColumn, TextColumn, TimeElapsedColumn

from Artefact.error_handler import handle_error, ValidationError, with_error_handling

console = Console()
logger = logging.getLogger(__name__)

# Try to import optional dependencies
try:
    from PIL import Image
    from PIL.ExifTags import TAGS, GPSTAGS
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    logger.debug("PIL not available - image metadata extraction limited")

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    logger.debug("PyPDF2 not available - PDF metadata extraction limited")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.debug("python-docx not available - DOCX metadata extraction limited")

try:
    import mutagen
    MUTAGEN_AVAILABLE = True
except ImportError:
    MUTAGEN_AVAILABLE = False
    logger.debug("mutagen not available - media metadata extraction limited")

try:
    import pefile
    PEFILE_AVAILABLE = True
except ImportError:
    PEFILE_AVAILABLE = False
    logger.debug("pefile not available - PE metadata extraction limited")

try:
    import elftools.elf.elffile as elftools
    PYELFTOOLS_AVAILABLE = True
except ImportError:
    PYELFTOOLS_AVAILABLE = False
    logger.debug("pyelftools not available - ELF metadata extraction limited")


@with_error_handling("extract_metadata")
def extract_metadata(file_path: Path, deep: bool = False, include_exif: bool = True) -> Dict[str, Any]:
    """
    Extract metadata from various file types.
    
    Args:
        file_path: Path to the file
        deep: Use external tools (exiftool) for deep extraction
        include_exif: Include EXIF data for images
        
    Returns:
        Dictionary containing metadata including timestamps
        
    Raises:
        ValidationError: If file doesn't exist
    """
    if not isinstance(file_path, Path):
        file_path = Path(file_path)
    
    if not file_path.exists():
        logger.warning(f"File does not exist: {file_path}")
        return {"timestamps": [], "error": f"File not found: {file_path}"}
    
    if not file_path.is_file():
        logger.warning(f"Path is not a file: {file_path}")
        return {"timestamps": [], "error": f"Path is not a file: {file_path}"}
    
    result = {
        "file_path": str(file_path),
        "file_size": file_path.stat().st_size,
        "timestamps": [],
        "metadata": {},
        "extractor": "artefact"
    }
    
    # Add basic file system timestamps
    try:
        stat = file_path.stat()
        result["timestamps"].extend([
            {
                "label": "File Created",
                "value": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "source": "filesystem"
            },
            {
                "label": "File Modified", 
                "value": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "source": "filesystem"
            },
            {
                "label": "File Accessed",
                "value": datetime.fromtimestamp(stat.st_atime).isoformat(), 
                "source": "filesystem"
            }
        ])
    except Exception as e:
        logger.warning(f"Failed to get filesystem timestamps: {e}")
    
    # Determine file type and extract accordingly
    file_ext = file_path.suffix.lower()
    
    # First try magic number detection
    magic_number = None
    try:
        with open(file_path, 'rb') as f:
            magic_number = f.read(4)
    except Exception:
        pass
    
    if deep:
        # Use exiftool for comprehensive extraction
        _extract_with_exiftool(file_path, result)
    
    # Check binary formats first by magic numbers
    if magic_number == b'MZ\x90\x00':
        # PE file
        _extract_pe_metadata(file_path, result)
    elif magic_number == b'\x7fELF':
        # ELF file
        _extract_elf_metadata(file_path, result)
    
    # Then check by extension
    if file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp'] and include_exif:
        _extract_image_metadata(file_path, result)
    elif file_ext == '.pdf':
        _extract_pdf_metadata(file_path, result)
    elif file_ext in ['.doc', '.docx']:
        _extract_document_metadata(file_path, result)
    elif file_ext in ['.mp4', '.avi', '.mov', '.mp3', '.wav', '.ogg', '.flac', '.m4a', '.wma']:
        _extract_media_metadata(file_path, result)
    elif not (magic_number == b'MZ\x90\x00' or magic_number == b'\x7fELF'):
        # Only log if we haven't already identified it as a binary
        logger.info(f"No specific metadata extractor for {file_ext} files")
    
    return result


def _extract_image_metadata(file_path: Path, result: Dict[str, Any]):
    """Extract metadata from image files using PIL."""
    if not PIL_AVAILABLE:
        logger.warning("PIL not available - skipping image metadata extraction")
        return
    
    try:
        with Image.open(file_path) as img:
            # Basic image info
            result["metadata"]["image"] = {
                "format": img.format,
                "mode": img.mode,
                "size": img.size,
                "has_transparency": img.mode in ('RGBA', 'LA') or 'transparency' in img.info
            }
            
            # Extract EXIF data
            if hasattr(img, '_getexif') and img._getexif():
                exif_data = img._getexif()
                exif_dict = {}
                
                for tag, value in exif_data.items():
                    tag_name = TAGS.get(tag, tag)
                    
                    # Handle special EXIF tags
                    if tag_name == 'GPSInfo':
                        gps_dict = {}
                        for gps_tag, gps_value in value.items():
                            gps_tag_name = GPSTAGS.get(gps_tag, gps_tag)
                            gps_dict[gps_tag_name] = gps_value
                        exif_dict[tag_name] = gps_dict
                    elif tag_name in ['DateTime', 'DateTimeOriginal', 'DateTimeDigitized']:
                        # Add timestamp
                        try:
                            dt = datetime.strptime(str(value), '%Y:%m:%d %H:%M:%S')
                            result["timestamps"].append({
                                "label": f"EXIF {tag_name}",
                                "value": dt.isoformat(),
                                "source": "exif"
                            })
                            exif_dict[tag_name] = str(value)
                        except ValueError:
                            exif_dict[tag_name] = str(value)
                    else:
                        exif_dict[tag_name] = str(value) if not isinstance(value, (int, float)) else value
                
                result["metadata"]["exif"] = exif_dict
                
    except Exception as e:
        logger.warning(f"Failed to extract image metadata from {file_path}: {e}")


def _extract_pdf_metadata(file_path: Path, result: Dict[str, Any]):
    """Extract metadata from PDF files."""
    if not PYPDF2_AVAILABLE:
        logger.warning("PyPDF2 not available - skipping PDF metadata extraction")
        return
    
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            
            # Basic PDF info
            result["metadata"]["pdf"] = {
                "pages": len(reader.pages),
                "encrypted": reader.is_encrypted
            }
            
            # Extract document metadata
            if reader.metadata:
                pdf_metadata = {}
                
                for key, value in reader.metadata.items():
                    clean_key = key.replace('/', '') if key.startswith('/') else key
                    pdf_metadata[clean_key] = str(value)
                    
                    # Check for timestamp fields
                    if any(date_field in clean_key.lower() for date_field in 
                          ['creationdate', 'moddate', 'date']):
                        try:
                            # PDF dates are often in format: D:YYYYMMDDHHmmSS
                            date_str = str(value)
                            if date_str.startswith('D:'):
                                date_str = date_str[2:]
                            
                            # Try to parse various date formats
                            for fmt in ['%Y%m%d%H%M%S', '%Y%m%d', '%Y-%m-%d %H:%M:%S']:
                                try:
                                    dt = datetime.strptime(date_str[:len(fmt.replace('%', ''))], fmt)
                                    result["timestamps"].append({
                                        "label": f"PDF {clean_key}",
                                        "value": dt.isoformat(),
                                        "source": "pdf"
                                    })
                                    break
                                except ValueError:
                                    continue
                        except Exception as e:
                            logger.debug(f"Failed to parse PDF date {value}: {e}")
                
                result["metadata"]["pdf_metadata"] = pdf_metadata
                
    except Exception as e:
        logger.warning(f"Failed to extract PDF metadata from {file_path}: {e}")


def _extract_document_metadata(file_path: Path, result: Dict[str, Any]):
    """Extract metadata from document files."""
    if not DOCX_AVAILABLE:
        logger.warning("python-docx not available - limited document metadata extraction")
        result["metadata"]["document"] = {
            "type": "office_document",
            "extension": file_path.suffix
        }
        return
        
    try:
        # Extract DOCX metadata
        if file_path.suffix.lower() == '.docx':
            doc = DocxDocument(file_path)
            core_props = doc.core_properties
            
            doc_metadata = {
                "title": core_props.title,
                "author": core_props.author,
                "company": core_props.company,
                "category": core_props.category,
                "comments": core_props.comments,
                "content_status": core_props.content_status,
                "subject": core_props.subject,
                "version": core_props.version,
                "language": core_props.language,
                "keywords": core_props.keywords,
                "last_modified_by": core_props.last_modified_by,
                "last_printed": core_props.last_printed,
                "revision": core_props.revision,
                "document_statistics": {
                    "paragraphs": len(doc.paragraphs),
                    "sections": len(doc.sections),
                    "tables": len(doc.tables)
                }
            }
            
            # Add timestamps
            for field, value in [
                ("created", core_props.created),
                ("modified", core_props.modified),
                ("last_printed", core_props.last_printed)
            ]:
                if value:
                    try:
                        if isinstance(value, str):
                            dt = datetime.fromisoformat(value.replace('Z', '+00:00'))
                        else:
                            dt = value.replace(tzinfo=timezone.utc)
                            
                        result["timestamps"].append({
                            "label": f"Document {field}",
                            "value": dt.isoformat(),
                            "source": "docx_metadata"
                        })
                    except Exception as e:
                        logger.debug(f"Failed to parse document timestamp {field}: {e}")
            
            result["metadata"]["document"] = doc_metadata
            
        else:
            result["metadata"]["document"] = {
                "type": "office_document",
                "extension": file_path.suffix,
                "note": "Metadata extraction not implemented for this format"
            }
        
    except Exception as e:
        logger.warning(f"Failed to extract document metadata from {file_path}: {e}")
        result["metadata"]["document"] = {
            "type": "office_document",
            "extension": file_path.suffix,
            "error": str(e)
        }


def _extract_media_metadata(file_path: Path, result: Dict[str, Any]):
    """Extract metadata from media files."""
    if not MUTAGEN_AVAILABLE:
        logger.warning("mutagen not available - limited media metadata extraction")
        result["metadata"]["media"] = {
            "type": "media_file",
            "extension": file_path.suffix
        }
        return
        
    try:
        # Load file with mutagen
        media_file = mutagen.File(file_path)
        
        if media_file is None:
            result["metadata"]["media"] = {
                "type": "media_file",
                "extension": file_path.suffix,
                "note": "Unsupported media format"
            }
            return
            
        # Extract common metadata
        media_metadata = {
            "type": "media_file",
            "extension": file_path.suffix,
            "format": media_file.mime[0] if hasattr(media_file, 'mime') else None,
            "length": media_file.info.length if hasattr(media_file.info, 'length') else None,
            "bitrate": media_file.info.bitrate if hasattr(media_file.info, 'bitrate') else None,
            "channels": media_file.info.channels if hasattr(media_file.info, 'channels') else None,
            "sample_rate": media_file.info.sample_rate if hasattr(media_file.info, 'sample_rate') else None,
            "bits_per_sample": media_file.info.bits_per_sample if hasattr(media_file.info, 'bits_per_sample') else None
        }
        
        # Extract tags
        if hasattr(media_file, 'tags'):
            media_metadata["tags"] = {}
            for key, value in media_file.tags.items():
                # Clean up tag names
                clean_key = key.lower().replace(':', '_').replace('/', '_')
                if isinstance(value, (list, tuple)):
                    value = value[0] if value else None
                media_metadata["tags"][clean_key] = str(value)
                
                # Check for date/time tags
                if any(date_field in clean_key.lower() for date_field in ['date', 'time', 'created', 'modified']):
                    try:
                        # Try common date formats
                        for fmt in ['%Y-%m-%d', '%Y:%m:%d', '%Y-%m-%d %H:%M:%S', '%Y:%m:%d %H:%M:%S']:
                            try:
                                dt = datetime.strptime(str(value)[:19], fmt)
                                result["timestamps"].append({
                                    "label": f"Media {clean_key}",
                                    "value": dt.isoformat(),
                                    "source": "media_metadata"
                                })
                                break
                            except ValueError:
                                continue
                    except Exception as e:
                        logger.debug(f"Failed to parse media timestamp {clean_key}: {e}")
        
        result["metadata"]["media"] = media_metadata
        
    except Exception as e:
        logger.warning(f"Failed to extract media metadata from {file_path}: {e}")
        result["metadata"]["media"] = {
            "type": "media_file",
            "extension": file_path.suffix,
            "error": str(e)
        }


def _extract_with_exiftool(file_path: Path, result: Dict[str, Any]):
    """Extract metadata using exiftool (if available)."""
    try:
        # Try to run exiftool
        cmd = ['exiftool', '-json', '-coordFormat', '%.6f', str(file_path)]
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        
        if proc.returncode == 0:
            exif_data = json.loads(proc.stdout)[0]
            result["metadata"]["exiftool"] = exif_data
            result["extractor"] = "exiftool"
            
            # Extract timestamps from exiftool output
            for key, value in exif_data.items():
                if any(date_field in key.lower() for date_field in 
                      ['date', 'time', 'created', 'modified']):
                    try:
                        # Try to parse the timestamp
                        for fmt in [
                            '%Y:%m:%d %H:%M:%S',
                            '%Y-%m-%d %H:%M:%S',
                            '%Y:%m:%d %H:%M:%S%z',
                            '%Y-%m-%dT%H:%M:%S'
                        ]:
                            try:
                                dt = datetime.strptime(str(value)[:19], fmt)
                                result["timestamps"].append({
                                    "label": f"ExifTool {key}",
                                    "value": dt.isoformat(),
                                    "source": "exiftool"
                                })
                                break
                            except ValueError:
                                continue
                    except Exception as e:
                        logger.debug(f"Failed to parse exiftool timestamp {key}={value}: {e}")
        else:
            logger.warning(f"exiftool failed: {proc.stderr}")
            
    except FileNotFoundError:
        logger.debug("exiftool not found - falling back to basic extraction")
    except subprocess.TimeoutExpired:
        logger.warning("exiftool timed out")
    except Exception as e:
        logger.warning(f"Failed to run exiftool: {e}")


def display_metadata(metadata: Dict[str, Any], show_timestamps: bool = True):
    """Display metadata in a formatted table."""
    if not metadata:
        console.print("[yellow]No metadata available[/]")
        return
    
    # File info table
    if "file_path" in metadata:
        file_table = Table(title="File Information", show_lines=True)
        file_table.add_column("Property", style="cyan")
        file_table.add_column("Value", style="white")
        
        file_table.add_row("File Path", metadata["file_path"])
        if "file_size" in metadata:
            size_str = _format_file_size(metadata["file_size"])
            file_table.add_row("File Size", size_str)
        if "extractor" in metadata:
            file_table.add_row("Extractor", metadata["extractor"])
        
        console.print(file_table)
    
    # Timestamps table
    if show_timestamps and metadata.get("timestamps"):
        ts_table = Table(title="Timestamps", show_lines=True)
        ts_table.add_column("Source", style="cyan")
        ts_table.add_column("Label", style="yellow")
        ts_table.add_column("Timestamp", style="green")
        
        for ts in metadata["timestamps"]:
            ts_table.add_row(
                ts.get("source", "unknown"),
                ts.get("label", "Unknown"),
                ts.get("value", "N/A")
            )
        
        console.print(ts_table)
    
    # Metadata details
    if metadata.get("metadata"):
        for section, data in metadata["metadata"].items():
            if isinstance(data, dict) and data:
                meta_table = Table(title=f"{section.title()} Metadata", show_lines=True)
                meta_table.add_column("Property", style="cyan")
                meta_table.add_column("Value", style="white")
                
                for key, value in data.items():
                    # Truncate very long values
                    str_value = str(value)
                    if len(str_value) > 100:
                        str_value = str_value[:97] + "..."
                    meta_table.add_row(key, str_value)
                
                console.print(meta_table)


def _extract_pe_metadata(file_path: Path, result: Dict[str, Any]):
    """Extract metadata from PE files."""
    if not PEFILE_AVAILABLE:
        logger.warning("pefile not available - limited PE metadata extraction")
        return
    
    try:
        pe = pefile.PE(file_path)
        
        pe_metadata = {
            "type": "pe_file",
            "machine": pe.FILE_HEADER.Machine,
            "timestamp": datetime.fromtimestamp(pe.FILE_HEADER.TimeDateStamp).isoformat(),
            "subsystem": pe.OPTIONAL_HEADER.Subsystem,
            "dll": pe.is_dll(),
            "exe": pe.is_exe(),
            "driver": pe.is_driver(),
            "architecture": {
                332: "PE32",
                512: "PE32+",
                267: "ROM"
            }.get(pe.OPTIONAL_HEADER.Magic, "Unknown"),
            "os_version": f"{pe.OPTIONAL_HEADER.MajorOperatingSystemVersion}.{pe.OPTIONAL_HEADER.MinorOperatingSystemVersion}",
            "image_version": f"{pe.OPTIONAL_HEADER.MajorImageVersion}.{pe.OPTIONAL_HEADER.MinorImageVersion}",
            "sections": []
        }
        
        # Get sections
        for section in pe.sections:
            pe_metadata["sections"].append({
                "name": section.Name.decode().rstrip('\x00'),
                "size": section.SizeOfRawData,
                "virtual_address": section.VirtualAddress,
                "characteristics": hex(section.Characteristics)
            })
        
        # Get imports
        if hasattr(pe, 'DIRECTORY_ENTRY_IMPORT'):
            pe_metadata["imports"] = []
            for entry in pe.DIRECTORY_ENTRY_IMPORT:
                dll_imports = {
                    "dll": entry.dll.decode() if hasattr(entry, 'dll') else "unknown",
                    "functions": []
                }
                if hasattr(entry, 'imports'):
                    for imp in entry.imports[:100]:  # Limit to first 100
                        if hasattr(imp, 'name') and imp.name:
                            dll_imports["functions"].append(imp.name.decode())
                pe_metadata["imports"].append(dll_imports)
        
        # Get exports
        if hasattr(pe, 'DIRECTORY_ENTRY_EXPORT'):
            pe_metadata["exports"] = []
            if hasattr(pe.DIRECTORY_ENTRY_EXPORT, 'symbols'):
                for exp in pe.DIRECTORY_ENTRY_EXPORT.symbols[:100]:  # Limit to first 100
                    if hasattr(exp, 'name') and exp.name:
                        pe_metadata["exports"].append(exp.name.decode())
        
        # Add timestamp
        result["timestamps"].append({
            "label": "PE Compiled",
            "value": pe_metadata["timestamp"],
            "source": "pe_metadata"
        })
        
        result["metadata"]["pe"] = pe_metadata
        pe.close()
        
    except Exception as e:
        logger.warning(f"Failed to extract PE metadata from {file_path}: {e}")


def _extract_elf_metadata(file_path: Path, result: Dict[str, Any]):
    """Extract metadata from ELF files."""
    if not PYELFTOOLS_AVAILABLE:
        logger.warning("pyelftools not available - limited ELF metadata extraction")
        return
    
    try:
        with open(file_path, 'rb') as f:
            elf = elftools.ELFFile(f)
            
            elf_metadata = {
                "type": "elf_file",
                "machine": elf.get_machine_arch(),
                "elfclass": elf.elfclass,
                "data": elf.little_endian and "little endian" or "big endian",
                "type": elf.header['e_type'],
                "entry_point": hex(elf.header['e_entry']),
                "sections": [],
                "segments": [],
                "dynamic": {},
                "symbols": []
            }
            
            # Get sections
            for section in elf.iter_sections():
                sect_data = {
                    "name": section.name,
                    "type": section['sh_type'],
                    "flags": section['sh_flags'],
                    "size": section['sh_size']
                }
                elf_metadata["sections"].append(sect_data)
            
            # Get segments
            for segment in elf.iter_segments():
                seg_data = {
                    "type": segment['p_type'],
                    "flags": segment['p_flags'],
                    "size": segment['p_filesz']
                }
                elf_metadata["segments"].append(seg_data)
            
            # Get dynamic entries
            if elf.header['e_type'] == 'ET_DYN':
                dynamic = elf.get_section_by_name('.dynamic')
                if dynamic:
                    for tag in dynamic.iter_tags():
                        if tag.entry.d_tag == 'DT_NEEDED':
                            elf_metadata["dynamic"].setdefault("needed", []).append(tag.needed)
                        elif tag.entry.d_tag == 'DT_RPATH':
                            elf_metadata["dynamic"]["rpath"] = tag.rpath
                        elif tag.entry.d_tag == 'DT_RUNPATH':
                            elf_metadata["dynamic"]["runpath"] = tag.runpath
            
            # Get symbols
            symtab = elf.get_section_by_name('.symtab')
            if symtab:
                for sym in symtab.iter_symbols()[:100]:  # Limit to first 100
                    if sym.name:
                        sym_data = {
                            "name": sym.name,
                            "type": sym["st_info"]["type"],
                            "bind": sym["st_info"]["bind"],
                            "value": hex(sym["st_value"])
                        }
                        elf_metadata["symbols"].append(sym_data)
            
            result["metadata"]["elf"] = elf_metadata
            
    except Exception as e:
        logger.warning(f"Failed to extract ELF metadata from {file_path}: {e}")


def _format_file_size(size_bytes: int) -> str:
    """Format file size in human-readable format."""
    for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
        if size_bytes < 1024.0:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024.0
    return f"{size_bytes:.1f} PB"


def batch_extract_metadata(file_paths: List[Path], deep: bool = False) -> Dict[str, Dict[str, Any]]:
    """Extract metadata from multiple files."""
    results = {}
    
    for file_path in file_paths:
        try:
            results[str(file_path)] = extract_metadata(file_path, deep=deep)
        except Exception as e:
            logger.error(f"Failed to extract metadata from {file_path}: {e}")
            results[str(file_path)] = {"error": str(e), "timestamps": []}
    
    return results


if __name__ == "__main__":
    # CLI interface for standalone usage
    import argparse
    
    parser = argparse.ArgumentParser(description="Metadata Extraction Tool")
    parser.add_argument("-f", "--file", required=True, help="File to extract metadata from")
    parser.add_argument("--deep", action="store_true", help="Use exiftool for deeper extraction")
    parser.add_argument("--no-exif", action="store_true", help="Skip EXIF data extraction")
    parser.add_argument("--json", action="store_true", help="Output as JSON")
    parser.add_argument("--timestamps-only", action="store_true", help="Show only timestamps")
    
    args = parser.parse_args()
    
    file_path = Path(args.file)
    
    try:
        metadata = extract_metadata(
            file_path, 
            deep=args.deep, 
            include_exif=not args.no_exif
        )
        
        if args.json:
            console.print_json(json.dumps(metadata, indent=2, default=str))
        else:
            display_metadata(metadata, show_timestamps=not args.timestamps_only)
            
        if args.timestamps_only:
            console.print(f"\n[green]Found {len(metadata.get('timestamps', []))} timestamps[/]")
            
    except Exception as e:
        console.print(f"[red]Error:[/] {str(e)}")
        exit(1)
